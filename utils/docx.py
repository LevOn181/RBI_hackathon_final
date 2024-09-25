from docx import Document
from utils import img as img_extractor

def extract_text_from_docx(wrapper, docx_path):
        text = ""
        doc = Document(docx_path)

        # Extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"

        # Extract text from tables
        processed_cells = set()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell in processed_cells:
                        continue
                    processed_cells.add(cell)
                    text += cell.text + "\n"

        # Extract text from headers and footers
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                text += para.text + "\n"
            footer = section.footer
            for para in footer.paragraphs:
                text += para.text + "\n"

        # Extract images from document
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                images.append(rel.target_part.blob)

        # Extract text from images
        images_text = img_extractor.get_content_from_images_with_gpt(wrapper, images)

        print(images_text)
        return text + "\nText content from images: " + images_text